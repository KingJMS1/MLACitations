from copy import deepcopy

from appJar import gui
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from Citation import Citation

# auth = "Mahlon Scott"
# title = "Fun and more fun"
# containter = "In the Wild"
# publisher = "Mahlon Scott"
# pubdate = "12 May 2018"
# accessdate = "13 June 2018"
# url = "http://www.google.com"
# document = Document()
# a = Citation(author=auth, title=title, container=containter, publisher=publisher, pubdate=pubdate, accessdate=accessdate, url=url)
# a.write(document)
# document.save("Use.docx")


nSaved = 0
document = Document()
pformat = document.styles["Normal"].paragraph_format
pformat.first_line_indent = Pt(-36)
pformat.left_indent = Pt(36)
citlist = []

def saveCitation(button):
    global box
    if box.curselection() == ():
        kwargs = {}
        auth = app.getEntry("Author:")
        tit = app.getEntry("Title:")
        cont = app.getEntry("Container:")
        ver = app.getEntry("Version:")
        vol = app.getEntry("Volume, Number:")
        pub = app.getEntry("Publisher:")
        pdate = app.getEntry("Date Published:")
        loc = app.getEntry("Location:")
        adate = app.getEntry("Date Accessed:")
        url = app.getEntry("URL:")
        global nSaved
        if auth != "":
            kwargs["author"] = auth
        if tit != "":
            kwargs["title"] = tit
        if cont != "":
            kwargs["container"] = cont
        if ver != "":
            kwargs["version"] = ver
        if vol != "":
            kwargs["num"] = vol.split(", ")
        if pub != "":
            kwargs["publisher"] = pub
        if pdate != "":
            kwargs["pubdate"] = pdate
        if loc != "":
            kwargs["location"] = loc
        if adate != "":
            kwargs["accessdate"] = adate
        if url != "":
            kwargs["url"] = url
        try:
            a = Citation(**kwargs)
        except IndexError:
            app.popUp(title="Failure", message="Unable to save citation.")
            return
        app.clearAllEntries()
        app.setFocus("Author:")
        citlist.append(a)
        app.addListItem("Citations", a.getFirst())
        box.selection_clear(0)
        nSaved += 1
        app.popUp(title="Progress", message=str(nSaved) + " Citations Saved.")
    else:
        kwargs = {}
        auth = app.getEntry("Author:")
        tit = app.getEntry("Title:")
        cont = app.getEntry("Container:")
        ver = app.getEntry("Version:")
        vol = app.getEntry("Volume, Number:")
        pub = app.getEntry("Publisher:")
        pdate = app.getEntry("Date Published:")
        loc = app.getEntry("Location:")
        adate = app.getEntry("Date Accessed:")
        url = app.getEntry("URL:")
        if auth != "":
            kwargs["author"] = auth
        if tit != "":
            kwargs["title"] = tit
        if cont != "":
            kwargs["container"] = cont
        if ver != "":
            kwargs["version"] = ver
        if vol != "":
            kwargs["num"] = vol.split(", ")
        if pub != "":
            kwargs["publisher"] = pub
        if pdate != "":
            kwargs["pubdate"] = pdate
        if loc != "":
            kwargs["location"] = loc
        if adate != "":
            kwargs["accessdate"] = adate
        if url != "":
            kwargs["url"] = url
        try:
            a = Citation(**kwargs)
        except IndexError:
            app.popUp(title="Failure", message="Unable to save citation.")
            return
        app.clearAllEntries()
        app.setFocus("Author:")
        index = deepcopy(box.curselection()[0])
        box.selection_clear(0)
        app.removeListItem("Citations", citlist[index].getFirst())
        del citlist[index]
        citlist.append(a)
        app.addListItem("Citations", a.getFirst())
        box.selection_clear(0)
        app.popUp(title="Progress", message=str(nSaved) + "Total Citations Saved.")

def makeCitations(button):
    citlist.sort(key=Citation.getFirst)
    header = document.add_paragraph()
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    headerText = header.add_run("Works Cited")
    headerText.font.name = "Times New Roman"
    headerText.font.size = Pt(12)
    headerText.bold = True
    for e in citlist:
        e.write(document)
    document.save("Works Cited.docx")
    app.popUp(title="Progress", message="Citations Written!")


def load(onSelect):
    global box
    app.clearAllEntries()
    if box.curselection() == ():
        return
    citIndex = int(box.curselection()[0])
    cit = citlist[citIndex]
    for key in cit.keys:
        if key in ["num", "pubdate", "accessdate", "url"]:
            if key == "num":
                app.setEntry("Volume, Number:", cit.data[key][0] + ", " + cit.data[key][1])
            elif key == "pubdate":
                app.setEntry("Date Published:", cit.data[key])
            elif key == "accessdate":
                app.setEntry("Date Accessed:", cit.data[key])
            elif key == "url":
                app.setEntry("URL:", cit.data[key])
        else:
            app.setEntry(key.capitalize() + ":", cit.data[key])

def boxClear(button):
    global box
    box.selection_clear(0)
    app.clearAllEntries()


def delCitation(button):
    global box
    global nSaved
    index = deepcopy(box.curselection()[0])
    box.selection_clear(0)
    app.removeListItem("Citations", citlist[index].getFirst())
    del citlist[index]
    nSaved -= 1
    app.clearAllEntries()


def lol(button):
    print("Trigger")

app = gui("MLA8 Citation Generator")
app.addLabelEntry("Author:")
app.addLabelEntry("Title:")
app.addLabelEntry("Container:")
# app.addLabelEntry("Editors:")
# app.addLabelEntry("Translators:")
app.addLabelEntry("Version:")
app.addLabelEntry("Volume, Number:")
app.addLabelEntry("Publisher:")
app.addLabelEntry("Date Published:")
app.addLabelEntry("Location:")
app.addLabelEntry("Date Accessed:")
app.addLabelEntry("URL:")

app.addButton("Save Citation", saveCitation)
app.addButton("Make Works Cited", makeCitations)
app.setFocus("Author:")
app.addVerticalSeparator(0, 1, 0, 12)
app.addLabel("Saved Citations", row=0, column=2)

box = app.addListBox("Citations", column=2, rowspan=10)
box.bind("<<ListboxSelect>>", load)
box.configure(selectmode="single")
box.bind("<FocusIn>", load)

app.addButton("Clear Selection", boxClear, column=2, row=11)
app.addButton("Delete Citation", delCitation, column=2, row=12)

app.go()